"""
Aurelian Manufacturing — Markdown to DOCX Converter
Converts quality system Markdown documents to styled Word documents
per the Aurelian Manufacturing Design Manual.

Usage:
    python md_to_docx.py <input.md> <output.docx>
    python md_to_docx.py --all   (converts all three review documents)
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Aurelian Design Manual Colors ───
RED = RGBColor(0xF5, 0x05, 0x37)        # #F50537
BLACK = RGBColor(0x2B, 0x2B, 0x2B)      # #2B2B2B
GRAY = RGBColor(0x6B, 0x6B, 0x6B)       # #6B6B6B
PANEL_GRAY = RGBColor(0xF5, 0xF5, 0xF5) # #F5F5F5
DIVIDER = RGBColor(0xD3, 0xD3, 0xD3)    # #D3D3D3
WHITE = RGBColor(0xFF, 0xFF, 0xFF)       # #FFFFFF
HEADER_BG = "2B2B2B"
ALT_ROW_BG = "F5F5F5"
RED_HEX = "F50537"


def setup_styles(doc):
    """Configure document styles per Aurelian Design Manual §4.2"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = BLACK
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Title
    if 'Title' in doc.styles:
        s = doc.styles['Title']
        s.font.name = 'Calibri'
        s.font.size = Pt(28)
        s.font.bold = True
        s.font.color.rgb = BLACK
        s.paragraph_format.space_after = Pt(12)
        s.paragraph_format.space_before = Pt(0)

    # Heading 1
    s = doc.styles['Heading 1']
    s.font.name = 'Calibri'
    s.font.size = Pt(22)
    s.font.bold = True
    s.font.color.rgb = BLACK
    s.paragraph_format.space_before = Pt(24)
    s.paragraph_format.space_after = Pt(6)

    # Heading 2
    s = doc.styles['Heading 2']
    s.font.name = 'Calibri'
    s.font.size = Pt(16)
    s.font.bold = True
    s.font.color.rgb = RED
    s.paragraph_format.space_before = Pt(18)
    s.paragraph_format.space_after = Pt(6)

    # Heading 3
    s = doc.styles['Heading 3']
    s.font.name = 'Calibri'
    s.font.size = Pt(13)
    s.font.bold = True
    s.font.color.rgb = BLACK
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(4)

    # Heading 4
    s = doc.styles['Heading 4']
    s.font.name = 'Calibri'
    s.font.size = Pt(11)
    s.font.bold = True
    s.font.color.rgb = GRAY
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(4)


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a cell. Each border is (size_eighths, color_hex) or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(borders)

    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            sz, color = val
            el = parse_xml(
                f'<w:{side} {nsdecls("w")} w:val="single" w:sz="{sz}" '
                f'w:space="0" w:color="{color}"/>'
            )
            existing = borders.find(qn(f'w:{side}'))
            if existing is not None:
                borders.remove(existing)
            borders.append(el)


def remove_cell_borders(cell):
    """Remove all borders from a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(borders)


def add_table(doc, rows_data, has_header=True):
    """Add a styled table per Aurelian Design Manual §6."""
    if not rows_data or not rows_data[0]:
        return

    # Normalize: all rows must have the same number of columns
    num_cols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < num_cols:
            r.append('')

    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    # Remove default table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            # Clear default paragraph
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text).strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

            if i == 0 and has_header:
                # Header row
                set_cell_shading(cell, HEADER_BG)
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(11)
            else:
                run.font.color.rgb = BLACK
                # Alternating row shading
                if i % 2 == 0 and has_header:
                    set_cell_shading(cell, ALT_ROW_BG)

            # Cell padding
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            mar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'<w:top w:w="60" w:type="dxa"/>'
                f'<w:bottom w:w="60" w:type="dxa"/>'
                f'<w:left w:w="80" w:type="dxa"/>'
                f'<w:right w:w="80" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            existing_mar = tcPr.find(qn('w:tcMar'))
            if existing_mar is not None:
                tcPr.remove(existing_mar)
            tcPr.append(mar)

    doc.add_paragraph()  # spacing after table
    return table


def add_red_top_bar(doc):
    """Add a 3pt red decorative line at the top of the first page."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.space_before = Pt(0)
    # Add a red bottom border to simulate the bar
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="24" w:space="1" w:color="{RED_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_header_footer(doc, doc_title, date_str="2026-02-21"):
    """Add header and footer per Aurelian Design Manual §4.2."""
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Aurelian Manufacturing")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False

    # Create a 3-column table for footer layout
    ft = footer.add_table(1, 3, width=Inches(6.27))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders from footer table
    for cell in ft.rows[0].cells:
        remove_cell_borders(cell)

    # Left: CONFIDENTIAL
    left_cell = ft.rows[0].cells[0]
    left_cell.paragraphs[0].clear()
    run = left_cell.paragraphs[0].add_run("CONFIDENTIAL")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Center: Page number
    center_cell = ft.rows[0].cells[1]
    center_cell.paragraphs[0].clear()
    center_p = center_cell.paragraphs[0]
    center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = center_p.add_run()
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    # Add page number field
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE \\* MERGEFORMAT ">'
        f'<w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'<w:sz w:val="18"/><w:color w:val="6B6B6B"/></w:rPr>'
        f'<w:t>1</w:t></w:r></w:fldSimple>'
    )
    center_p._p.append(parse_xml(fld_xml))

    # Right: Date
    right_cell = ft.rows[0].cells[2]
    right_cell.paragraphs[0].clear()
    run = right_cell.paragraphs[0].add_run(date_str)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT


def parse_markdown_table(lines):
    """Parse markdown table lines into list of lists."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        # Skip separator rows (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    return rows


def clean_md_formatting(text):
    """Remove markdown inline formatting for plain text."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)       # italic
    text = re.sub(r'`(.+?)`', r'\1', text)         # code
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text) # links
    return text.strip()


def process_markdown(doc, md_text):
    """Parse markdown content and add to document with proper styling."""
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    in_list = False
    list_items = []
    is_checkbox_list = False

    while i < len(lines):
        line = lines[i]

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                # Add shading to paragraph
                pPr = p._p.get_or_add_pPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                pPr.append(shading)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = BLACK
                code_lines = []
                in_code_block = False
            else:
                # Flush any pending table
                if in_table:
                    rows = parse_markdown_table(table_lines)
                    if rows:
                        add_table(doc, rows)
                    table_lines = []
                    in_table = False
                if in_list:
                    in_list = False
                    list_items = []
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            if not in_table:
                # Flush list
                if in_list:
                    in_list = False
                    list_items = []
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            rows = parse_markdown_table(table_lines)
            if rows:
                add_table(doc, rows)
            table_lines = []
            in_table = False

        stripped = line.strip()

        # Empty line
        if not stripped:
            if in_list:
                in_list = False
                list_items = []
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="D3D3D3"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
        if heading_match:
            if in_list:
                in_list = False
                list_items = []
            level = len(heading_match.group(1))
            text = clean_md_formatting(heading_match.group(2))

            if level == 1:
                doc.add_heading(text, level=1)
            elif level == 2:
                doc.add_heading(text, level=2)
            elif level == 3:
                doc.add_heading(text, level=3)
            elif level >= 4:
                doc.add_heading(text, level=4)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = clean_md_formatting(stripped.lstrip('> '))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # Add left red border for callout style
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="24" w:space="4" w:color="{RED_HEX}"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            # Gray background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            pPr.append(shading)
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
            run.font.italic = True
            i += 1
            continue

        # Checkbox list items
        checkbox_match = re.match(r'^- \[[ x]\]\s+(.+)', stripped)
        if checkbox_match:
            text = clean_md_formatting(checkbox_match.group(1))
            is_checked = stripped.startswith('- [x]')
            prefix = "  [x] " if is_checked else "  [ ] "
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(prefix + text)
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^[-*]\s+(.+)', stripped)
        if bullet_match:
            text = clean_md_formatting(bullet_match.group(1))
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_formatted_run(p, text)
            i += 1
            continue

        # Numbered list
        numbered_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if numbered_match:
            text = clean_md_formatting(numbered_match.group(2))
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_formatted_run(p, text)
            i += 1
            continue

        # Indented sub-items (  - item)
        sub_bullet_match = re.match(r'^  [-*]\s+(.+)', stripped)
        if sub_bullet_match:
            text = clean_md_formatting(sub_bullet_match.group(1))
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(2.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            add_formatted_run(p, text)
            i += 1
            continue

        # Regular paragraph
        text = stripped
        p = doc.add_paragraph()
        add_formatted_run(p, text)
        i += 1

    # Flush remaining table
    if in_table and table_lines:
        rows = parse_markdown_table(table_lines)
        if rows:
            add_table(doc, rows)


def add_formatted_run(paragraph, text):
    """Add text to a paragraph with inline markdown formatting (bold, italic, code)."""
    # Clear existing runs
    for run in paragraph.runs:
        run.clear()

    # Simple approach: split on formatting markers
    # Pattern matches **bold**, *italic*, `code`, and plain text
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)

    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
        else:
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK


def convert_md_to_docx(md_path, docx_path, title=None, date_str="2026-02-21"):
    """Main conversion function."""
    print(f"Reading: {md_path}")
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Extract title from first heading if not provided
    if not title:
        title_match = re.match(r'^#\s+(.+)', md_text, re.MULTILINE)
        title = clean_md_formatting(title_match.group(1)) if title_match else "Document"

    doc = Document()

    # Page setup — A4, 2.5cm margins
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    # Setup styles
    setup_styles(doc)

    # Red top bar
    add_red_top_bar(doc)

    # Title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = tp.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = BLACK
    tp.paragraph_format.space_after = Pt(4)

    # Red accent bar under title
    bar = doc.add_paragraph()
    bar.paragraph_format.space_after = Pt(12)
    pPr = bar._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="{RED_HEX}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Header and footer
    add_header_footer(doc, title, date_str)

    # Skip the first heading in markdown (already used as title)
    md_text_no_title = re.sub(r'^#\s+.+\n?', '', md_text, count=1)

    # Process markdown content
    print(f"Processing content...")
    process_markdown(doc, md_text_no_title)

    # Save
    print(f"Saving: {docx_path}")
    doc.save(docx_path)
    print(f"Done: {docx_path}")


def main():
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    qs_dir = os.path.join(base_dir, "Quality System")

    if len(sys.argv) > 1 and sys.argv[1] == '--all':
        documents = [
            {
                'input': os.path.join(qs_dir, 'AM-MVP-001-MVP-Scope-Definition.md'),
                'output': os.path.join(qs_dir, 'AM-MVP-001-Review.docx'),
                'title': 'MVP SCOPE DEFINITION — DIGITAL MRB BUILDER PHASE 0',
            },
            {
                'input': os.path.join(qs_dir, 'AM-TS-001-IT-System-Technical-Specification.md'),
                'output': os.path.join(qs_dir, 'AM-TS-001-Review.docx'),
                'title': 'IT/SYSTEM TECHNICAL SPECIFICATION — DIGITAL MRB BUILDER',
            },
            {
                'input': os.path.join(qs_dir, 'AM-FM-001-Form-Templates.md'),
                'output': os.path.join(qs_dir, 'AM-FM-001-Review.docx'),
                'title': 'FORM TEMPLATE SPECIFICATIONS — DIGITAL MRB BUILDER',
            },
        ]
        for d in documents:
            if os.path.exists(d['input']):
                convert_md_to_docx(d['input'], d['output'], d['title'])
            else:
                print(f"WARNING: {d['input']} not found, skipping.")
        print("\n=== All documents generated ===")
    elif len(sys.argv) == 3:
        convert_md_to_docx(sys.argv[1], sys.argv[2])
    else:
        print("Usage:")
        print("  python md_to_docx.py --all")
        print("  python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)


if __name__ == '__main__':
    main()

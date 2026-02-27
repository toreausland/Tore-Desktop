"""
Tier 2 Fixes — Aurelian VDR Design Compliance
===============================================
Fixes 4 minor issues in files that already comply with Design Manual v2.0.

Issue 1: 02_Valuation_Methodology_V2.docx — Remove "REV6" from footer
Issue 2: 03_Market_Trends_Projections.docx — Fix header (03.05 → 03.06) + body refs
Issue 3: 02_Financing_Plan_V3.docx — Fix REV6 bare references in references table + body
Issue 4: 03_Market_Trends_Projections.docx — Add VDR References table at end

NOTE: 02_Economic_Tables_Projections_REV6.docx is NEVER touched. Master document.
"""

import sys
import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

VDR_ROOT = "Aurelian_VDR"
BACKUP_DIR = "_BACKUPS_TIER2"
LOG = []

def log(msg):
    LOG.append(msg)
    print(msg)

def backup(filepath):
    """Create backup before modifying."""
    os.makedirs(BACKUP_DIR, exist_ok=True)
    fname = os.path.basename(filepath)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = os.path.join(BACKUP_DIR, f"{ts}_{fname}")
    shutil.copy2(filepath, backup_path)
    log(f"  ✓ Backup: {backup_path}")
    return backup_path


# ============================================================
# ISSUE 1: Remove REV6 from footer of Valuation Methodology
# ============================================================
def fix_issue_1():
    log("\n" + "="*60)
    log("ISSUE 1: 02_Valuation_Methodology_V2.docx")
    log("Action: Remove 'REV6' from footer, keep date")
    log("="*60)

    fpath = os.path.join(VDR_ROOT, "02_Financial", "2.3_Valuation_Basis", "02_Valuation_Methodology_V2.docx")
    backup(fpath)

    doc = Document(fpath)
    changes = 0

    for section in doc.sections:
        footer = section.footer
        if footer:
            for para in footer.paragraphs:
                for run in para.runs:
                    if "REV6" in run.text:
                        old_text = run.text
                        # Replace "February 2026 — REV6" with "February 2026"
                        run.text = run.text.replace(" — REV6", "").replace("— REV6", "").replace("REV6", "")
                        log(f"  BEFORE: \"{old_text}\"")
                        log(f"  AFTER:  \"{run.text}\"")
                        changes += 1

    doc.save(fpath)
    log(f"  ✓ Saved. {changes} change(s) applied.")
    return changes


# ============================================================
# ISSUE 2: Fix header in Market Trends (03.05 → 03.06)
# Also fix body text references to 03.05
# ============================================================
def fix_issue_2():
    log("\n" + "="*60)
    log("ISSUE 2: 03_Market_Trends_Projections.docx")
    log("Action: Fix VDR reference 03.05 → 03.06 in header + body")
    log("="*60)

    fpath = os.path.join(VDR_ROOT, "03_Commercial_Market", "3.1_Market_Analysis", "03_Market_Trends_Projections.docx")
    backup(fpath)

    doc = Document(fpath)
    changes = 0

    # Fix headers
    for section in doc.sections:
        header = section.header
        if header:
            for para in header.paragraphs:
                for run in para.runs:
                    if "03.05" in run.text:
                        old_text = run.text
                        run.text = run.text.replace("03.05", "03.06")
                        log(f"  HEADER BEFORE: \"{old_text}\"")
                        log(f"  HEADER AFTER:  \"{run.text}\"")
                        changes += 1

    # Fix body text references to 03.05
    for para in doc.paragraphs:
        for run in para.runs:
            if "03.05" in run.text:
                old_text = run.text
                run.text = run.text.replace("03.05", "03.06")
                log(f"  BODY BEFORE: \"{old_text}\"")
                log(f"  BODY AFTER:  \"{run.text}\"")
                changes += 1

    doc.save(fpath)
    log(f"  ✓ Saved. {changes} change(s) applied.")
    return changes


# ============================================================
# ISSUE 3: Fix bare "REV6" references in Financing Plan
# The doc already has a References table — fix the entries
# ============================================================
def fix_issue_3():
    log("\n" + "="*60)
    log("ISSUE 3: 02_Financing_Plan_V3.docx")
    log("Action: Fix bare 'REV6' references in body text and references table")
    log("="*60)

    fpath = os.path.join(VDR_ROOT, "02_Financial", "2.2_Financing_Plan", "02_Financing_Plan_V3.docx")
    backup(fpath)

    doc = Document(fpath)
    changes = 0

    # Fix body text: replace bare "REV6" with proper reference
    # Patterns found:
    #   "REV6 (VDR 02.04)" → "02 Economic Tables & Projections (VDR 02.04)"
    #   "REV6 §1.2b" → "02 Economic Tables & Projections, §1.2b (VDR 02.04)"
    #   "REV6-projeksjoner" → heading fix
    for para in doc.paragraphs:
        for run in para.runs:
            if "REV6" in run.text:
                old_text = run.text
                new_text = run.text

                # Handle "REV6 (VDR 02.04)" — already has VDR ref
                new_text = new_text.replace(
                    "REV6 (VDR 02.04)",
                    "02 Economic Tables & Projections (VDR 02.04)"
                )

                # Handle "REV6 §" section references
                new_text = new_text.replace(
                    "REV6 §",
                    "02 Economic Tables & Projections, § "
                )

                # Handle heading "REV6-projeksjoner"
                new_text = new_text.replace(
                    "REV6-projeksjoner",
                    "02 Economic Tables & Projections"
                )

                # Catch any remaining bare "REV6" (shouldn't be any after above)
                # But be careful not to double-replace
                if "REV6" in new_text and "Economic Tables" not in new_text:
                    new_text = new_text.replace("REV6", "02 Economic Tables & Projections")

                if new_text != old_text:
                    run.text = new_text
                    log(f"  BEFORE: \"{old_text[:100]}\"")
                    log(f"  AFTER:  \"{new_text[:100]}\"")
                    changes += 1

    # Fix references table: row 1 says "REV6 (Economic Tables & Projections)"
    # Should be "02 Economic Tables & Projections (master document)"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "REV6" in run.text:
                            old_text = run.text
                            run.text = run.text.replace(
                                "REV6 (Economic Tables & Projections)",
                                "02 Economic Tables & Projections (master document)"
                            )
                            # Catch other patterns
                            if "REV6" in run.text:
                                run.text = run.text.replace("REV6", "02 Economic Tables & Projections")
                            if run.text != old_text:
                                log(f"  TABLE BEFORE: \"{old_text}\"")
                                log(f"  TABLE AFTER:  \"{run.text}\"")
                                changes += 1

    doc.save(fpath)
    log(f"  ✓ Saved. {changes} change(s) applied.")
    return changes


# ============================================================
# ISSUE 4: Add VDR References table to Market Trends
# ============================================================
def fix_issue_4():
    log("\n" + "="*60)
    log("ISSUE 4: 03_Market_Trends_Projections.docx")
    log("Action: Add VDR References table before contact info at end")
    log("="*60)

    fpath = os.path.join(VDR_ROOT, "03_Commercial_Market", "3.1_Market_Analysis", "03_Market_Trends_Projections.docx")
    # Already backed up in Issue 2 — but backup the Issue 2 result
    backup(fpath)

    doc = Document(fpath)

    # The document ends with:
    #   "Aurelian Manufacturing AS | VDR 03.06 | Confidential" (now fixed)
    #   "Contact: André Tandberg, CEO — ..."
    #   "Contact: Tore Ausland, VP ..."
    # We need to insert a References heading + table BEFORE those lines

    # Find the paragraph index for the closing contact block
    insert_before_idx = None
    for i, para in enumerate(doc.paragraphs):
        if "Aurelian Manufacturing AS" in para.text and "VDR 03.0" in para.text:
            insert_before_idx = i
            break

    if insert_before_idx is None:
        log("  ⚠ Could not find insertion point. Appending at end.")
        insert_before_idx = len(doc.paragraphs)

    # We need to work with the XML directly to insert before a specific paragraph
    # Get the element of the paragraph we want to insert before
    target_element = doc.paragraphs[insert_before_idx]._element

    # Create the References heading
    heading_para = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:pStyle w:val="Heading2"/></w:pPr>'
        f'  <w:r><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'    <w:b/><w:sz w:val="32"/><w:color w:val="F50537"/>'
        f'  </w:rPr><w:t>References</w:t></w:r>'
        f'</w:p>'
    )

    # Create spacer paragraph before references
    spacer = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:before="480"/></w:pPr>'
        f'</w:p>'
    )

    # Insert spacer + heading before the contact block
    target_element.addprevious(spacer)
    target_element.addprevious(heading_para)

    # Now create the references table
    # This document references VDR docs in its text
    ref_data = [
        ("VDR ref", "Document"),  # Header row
        ("02.04", "02 Economic Tables & Projections (master document)"),
        ("03.01", "Market Analysis (TAM/SAM/SOM)"),
        ("03.02", "CNC Benchmark & Competitive Landscape"),
    ]

    # Build table XML with Design Manual styling
    # Dark header (#2B2B2B), alternating rows, horizontal borders only
    tbl_xml = f'<w:tbl {nsdecls("w")}>'

    # Table properties
    tbl_xml += '''
    <w:tblPr>
      <w:tblStyle w:val="TableGrid"/>
      <w:tblW w:w="5000" w:type="pct"/>
      <w:tblBorders>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>
        <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      </w:tblBorders>
    </w:tblPr>
    <w:tblGrid>
      <w:gridCol w:w="2000"/>
      <w:gridCol w:w="7000"/>
    </w:tblGrid>
    '''

    for row_idx, (col1, col2) in enumerate(ref_data):
        if row_idx == 0:
            # Header row — dark background, white text
            tbl_xml += f'''
            <w:tr>
              <w:tc>
                <w:tcPr><w:shd w:val="clear" w:fill="2B2B2B"/></w:tcPr>
                <w:p><w:pPr><w:spacing w:before="60" w:after="60"/></w:pPr>
                  <w:r><w:rPr>
                    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                    <w:b/><w:sz w:val="22"/><w:color w:val="FFFFFF"/>
                  </w:rPr><w:t>{col1}</w:t></w:r>
                </w:p>
              </w:tc>
              <w:tc>
                <w:tcPr><w:shd w:val="clear" w:fill="2B2B2B"/></w:tcPr>
                <w:p><w:pPr><w:spacing w:before="60" w:after="60"/></w:pPr>
                  <w:r><w:rPr>
                    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                    <w:b/><w:sz w:val="22"/><w:color w:val="FFFFFF"/>
                  </w:rPr><w:t>{col2}</w:t></w:r>
                </w:p>
              </w:tc>
            </w:tr>'''
        else:
            # Body row — alternating fill
            fill = "F5F5F5" if row_idx % 2 == 0 else "FFFFFF"
            tbl_xml += f'''
            <w:tr>
              <w:tc>
                <w:tcPr><w:shd w:val="clear" w:fill="{fill}"/></w:tcPr>
                <w:p><w:pPr><w:spacing w:before="40" w:after="40"/></w:pPr>
                  <w:r><w:rPr>
                    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                    <w:sz w:val="22"/><w:color w:val="2B2B2B"/>
                  </w:rPr><w:t>{col1}</w:t></w:r>
                </w:p>
              </w:tc>
              <w:tc>
                <w:tcPr><w:shd w:val="clear" w:fill="{fill}"/></w:tcPr>
                <w:p><w:pPr><w:spacing w:before="40" w:after="40"/></w:pPr>
                  <w:r><w:rPr>
                    <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
                    <w:sz w:val="22"/><w:color w:val="2B2B2B"/>
                  </w:rPr><w:t>{col2}</w:t></w:r>
                </w:p>
              </w:tc>
            </w:tr>'''

    tbl_xml += '</w:tbl>'

    tbl_element = parse_xml(tbl_xml)
    target_element.addprevious(tbl_element)

    # Add spacer after table
    spacer2 = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr><w:spacing w:after="240"/></w:pPr>'
        f'</w:p>'
    )
    target_element.addprevious(spacer2)

    doc.save(fpath)
    log(f"  ✓ Saved. References heading + table inserted before contact block.")
    return 1


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    log("Aurelian VDR — Tier 2 Fixes")
    log(f"Date: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    log(f"Working directory: {os.getcwd()}")
    log("")

    total = 0
    total += fix_issue_1()
    total += fix_issue_2()
    total += fix_issue_3()
    total += fix_issue_4()

    log("\n" + "="*60)
    log(f"COMPLETE — {total} total change(s) across 3 files")
    log(f"Backups in: {BACKUP_DIR}/")
    log("="*60)

    # Summary
    log("\nSUMMARY OF CHANGES:")
    log("1. 02_Valuation_Methodology_V2.docx — footer: removed '— REV6'")
    log("2. 03_Market_Trends_Projections.docx — header + body: 03.05 → 03.06")
    log("3. 02_Financing_Plan_V3.docx — body + table: bare 'REV6' → full document name")
    log("4. 03_Market_Trends_Projections.docx — added VDR References table")
    log("\nNOTE: Master document (02_Economic_Tables_Projections_REV6) was NOT touched.")

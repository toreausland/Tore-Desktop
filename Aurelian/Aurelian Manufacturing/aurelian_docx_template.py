"""
Aurelian Manufacturing — Word Document Template Builder
=========================================================
Reusable python-docx functions for generating Design Manual v2.0 compliant documents.

Design Manual v2.0 compliance checklist:
  1. Calibri font throughout
  2. Aurelian Red #F50537 accent
  3. Header with "Aurelian Manufacturing"
  4. Footer with CONFIDENTIAL + page number + date
  5. Dark header tables (#2B2B2B)
  6. Alternating row shading (#F5F5F5)
  7. Horizontal borders only (#D3D3D3)
  8. References section at end
  9. No REV6, no banned names, no bare shortcuts
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# CONSTANTS — Design Manual v2.0
# ============================================================
AURELIAN_RED = RGBColor(0xF5, 0x05, 0x37)
NEAR_BLACK = RGBColor(0x2B, 0x2B, 0x2B)
SECONDARY = RGBColor(0x6B, 0x6B, 0x6B)
PANEL_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
DIVIDER_GRAY = RGBColor(0xD3, 0xD3, 0xD3)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PURE_BLACK = RGBColor(0x00, 0x00, 0x00)

FONT_NAME = "Calibri"

# Hex strings for XML
HEX_RED = "F50537"
HEX_NEAR_BLACK = "2B2B2B"
HEX_SECONDARY = "6B6B6B"
HEX_PANEL = "F5F5F5"
HEX_DIVIDER = "D3D3D3"
HEX_WHITE = "FFFFFF"


# ============================================================
# DOCUMENT CREATION
# ============================================================
def create_aurelian_doc():
    """Create a new Document with Aurelian page setup (A4, 2.5cm margins)."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    # Different first page header/footer
    section.different_first_page_header_footer = True

    return doc


# ============================================================
# STYLES — set up heading styles matching Design Manual
# ============================================================
def setup_styles(doc):
    """Configure all paragraph and heading styles per Design Manual v2.0."""
    styles = doc.styles

    # Normal (body text)
    normal = styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.font.color.rgb = NEAR_BLACK
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15

    # Title
    if 'Title' in [s.name for s in styles]:
        title = styles['Title']
        title.font.name = FONT_NAME
        title.font.size = Pt(28)
        title.font.bold = True
        title.font.color.rgb = NEAR_BLACK
        title.paragraph_format.space_after = Pt(12)
        title.paragraph_format.space_before = Pt(0)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = FONT_NAME
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = NEAR_BLACK
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = FONT_NAME
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = AURELIAN_RED
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = FONT_NAME
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = NEAR_BLACK
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # Heading 4
    h4 = styles['Heading 4']
    h4.font.name = FONT_NAME
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.color.rgb = SECONDARY
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)

    return doc


# ============================================================
# HEADER & FOOTER
# ============================================================
def setup_header(doc, doc_subtitle=""):
    """Add header: 'Aurelian Manufacturing' right-aligned in #6B6B6B with red underline."""
    section = doc.sections[0]

    # Main header (pages 2+)
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run("Aurelian Manufacturing")
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = SECONDARY

    if doc_subtitle:
        run2 = para.add_run(f"\t{doc_subtitle}")
        run2.font.name = FONT_NAME
        run2.font.size = Pt(9)
        run2.font.color.rgb = SECONDARY

    # Red underline below header
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{HEX_RED}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    return doc


def setup_footer(doc, date_str="February 2026"):
    """Add footer: CONFIDENTIAL left, page number center, date right."""
    section = doc.sections[0]

    # Main footer (pages 2+)
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Tab stops for center and right alignment
    pPr = para._element.get_or_add_pPr()
    tabs_xml = parse_xml(
        f'<w:tabs {nsdecls("w")}>'
        f'  <w:tab w:val="center" w:pos="4820"/>'
        f'  <w:tab w:val="right" w:pos="9639"/>'
        f'</w:tabs>'
    )
    pPr.append(tabs_xml)

    # CONFIDENTIAL (left)
    run_conf = para.add_run("CONFIDENTIAL")
    run_conf.font.name = FONT_NAME
    run_conf.font.size = Pt(9)
    run_conf.font.color.rgb = SECONDARY

    # Tab + Page number (center)
    run_tab1 = para.add_run("\t")
    run_page_label = para.add_run("Page ")
    run_page_label.font.name = FONT_NAME
    run_page_label.font.size = Pt(9)
    run_page_label.font.color.rgb = SECONDARY

    # Insert PAGE field
    fld_xml = (
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE ">'
        f'  <w:r><w:rPr>'
        f'    <w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="18"/><w:color w:val="{HEX_SECONDARY}"/>'
        f'  </w:rPr><w:t>2</w:t></w:r>'
        f'</w:fldSimple>'
    )
    para._element.append(parse_xml(fld_xml))

    # Tab + Date (right)
    run_tab2 = para.add_run("\t")
    run_date = para.add_run(date_str)
    run_date.font.name = FONT_NAME
    run_date.font.size = Pt(9)
    run_date.font.color.rgb = SECONDARY

    # Thin line above footer
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="1" w:color="{HEX_DIVIDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    return doc


# ============================================================
# TITLE PAGE
# ============================================================
def add_title_page(doc, title, subtitle="", vdr_ref="", date_str="February 2026"):
    """Add a cover/title page with Aurelian branding."""
    # Red decorative bar at top
    p_bar = doc.add_paragraph()
    p_bar.paragraph_format.space_before = Pt(0)
    p_bar.paragraph_format.space_after = Pt(48)
    pPr = p_bar._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="24" w:space="1" w:color="{HEX_RED}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Company name
    p_company = doc.add_paragraph()
    p_company.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_company.add_run("Aurelian Manufacturing")
    run.font.name = FONT_NAME
    run.font.size = Pt(14)
    run.font.color.rgb = SECONDARY
    p_company.paragraph_format.space_after = Pt(24)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(title)
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NEAR_BLACK
    p_title.paragraph_format.space_after = Pt(8)

    # Subtitle
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_sub.add_run(subtitle)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.font.color.rgb = SECONDARY
        p_sub.paragraph_format.space_after = Pt(24)

    # VDR reference
    if vdr_ref:
        p_vdr = doc.add_paragraph()
        run = p_vdr.add_run(vdr_ref)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = SECONDARY
        p_vdr.paragraph_format.space_after = Pt(4)

    # Date
    p_date = doc.add_paragraph()
    run = p_date.add_run(date_str)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = SECONDARY
    p_date.paragraph_format.space_after = Pt(36)

    # CONFIDENTIAL label
    p_conf = doc.add_paragraph()
    run = p_conf.add_run("CONFIDENTIAL")
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = SECONDARY
    run.font.bold = True
    p_conf.paragraph_format.space_after = Pt(0)

    # Page break after title page
    doc.add_page_break()

    return doc


# ============================================================
# TEXT HELPERS
# ============================================================
def add_heading(doc, text, level=1):
    """Add heading with correct style."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_NAME
    return h


def add_body(doc, text, bold=False, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    run.bold = bold
    run.italic = italic
    return p


def add_body_with_bold(doc, parts):
    """Add body paragraph with mixed bold/normal parts.
    parts = [("normal text", False), ("bold text", True), ("more normal", False)]
    """
    p = doc.add_paragraph()
    for text, is_bold in parts:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.color.rgb = NEAR_BLACK
        run.bold = is_bold
    return p


def add_bullet(doc, text, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = FONT_NAME
        run_b.font.size = Pt(11)
        run_b.font.color.rgb = NEAR_BLACK
        run_b.bold = True
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(11)
    run.font.color.rgb = NEAR_BLACK
    return p


def add_note(doc, text):
    """Add a note/caption paragraph in #6B6B6B, 9pt."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = SECONDARY
    run.italic = True
    return p


# ============================================================
# TABLE BUILDER — Design Manual v2.0
# ============================================================
def add_table(doc, headers, rows, col_widths=None):
    """Add a table with dark header row, alternating shading, horizontal borders only.

    headers: list of strings for header row
    rows: list of lists of strings for body rows
    col_widths: optional list of Cm values for column widths
    """
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Remove all borders first, then add horizontal only
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Set table width to 100%
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    # Borders: horizontal only
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_DIVIDER}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_DIVIDER}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_DIVIDER}"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    for j, header_text in enumerate(headers):
        cell = table.cell(0, j)
        # Dark background
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{HEX_NEAR_BLACK}"/>')
        cell._tc.get_or_add_tcPr().append(shd)
        # Text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header_text)
        run.font.name = FONT_NAME
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = WHITE
        # Cell padding
        _set_cell_padding(cell, top=60, bottom=60, left=80, right=80)

    # Body rows with alternating shading
    for i, row_data in enumerate(rows):
        fill = HEX_PANEL if i % 2 == 1 else HEX_WHITE
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i + 1, j)
            # Alternating fill
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>')
            cell._tc.get_or_add_tcPr().append(shd)
            # Text
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = NEAR_BLACK
            # Bold for first column if it looks like a label
            if j == 0 and len(row_data) > 2:
                pass  # Don't auto-bold, let caller decide
            # Cell padding
            _set_cell_padding(cell, top=40, bottom=40, left=80, right=80)

    # Set column widths if provided
    if col_widths:
        for j, width in enumerate(col_widths):
            for i in range(len(rows) + 1):
                cell = table.cell(i, j)
                cell.width = width

    # Space after table
    last_row_last_cell = table.cell(len(rows), num_cols - 1)

    return table


def add_table_with_bold_col(doc, headers, rows, bold_cols=None):
    """Same as add_table but allows specifying which columns should be bold.
    bold_cols: list of column indices to make bold (e.g., [0] for first column)
    """
    if bold_cols is None:
        bold_cols = []

    table = add_table(doc, headers, rows)

    # Apply bold to specified columns in body rows
    for i in range(len(rows)):
        for j in bold_cols:
            cell = table.cell(i + 1, j)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

    return table


def add_total_row_table(doc, headers, rows, total_row):
    """Add table with a bold total/summary row at the bottom with heavier top border.
    total_row: list of strings for the total row
    """
    all_rows = rows + [total_row]
    table = add_table(doc, headers, all_rows)

    # Bold the total row and add heavier top border
    total_idx = len(rows) + 1  # +1 for header
    num_cols = len(headers)
    for j in range(num_cols):
        cell = table.cell(total_idx, j)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        # Heavier top border on total row
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="{HEX_NEAR_BLACK}"/>'
            f'</w:tcBorders>'
        )
        tc_pr.append(tc_borders)

    return table


# ============================================================
# REFERENCES TABLE
# ============================================================
def add_references(doc, refs):
    """Add a References section per Design Manual v2.0 §13.5.

    refs: list of tuples (vdr_ref, document_title)
    e.g. [("02.04", "02 Economic Tables & Projections (master document)"),
          ("03.02", "CNC Benchmark & Competitive Landscape")]
    """
    add_heading(doc, "References", level=2)
    add_table(doc, ["VDR ref", "Document"], refs)
    return doc


# ============================================================
# CONTACT BLOCK (end of document)
# ============================================================
def add_contact_block(doc, vdr_ref=""):
    """Add contact block at end of document."""
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(24)
    parts = ["Aurelian Manufacturing AS"]
    if vdr_ref:
        parts.append(vdr_ref)
    parts.append("Confidential")
    run = p1.add_run("  |  ".join(parts))
    run.font.name = FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = SECONDARY

    for contact in [
        "Contact: Andre Tandberg, CEO — andre@aurelian.no",
        "Contact: Tore Ausland, VP Business Development — tore@aurelian.no"
    ]:
        p = doc.add_paragraph()
        run = p.add_run(contact)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.color.rgb = SECONDARY
        run.italic = True

    return doc


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def _set_cell_padding(cell, top=0, bottom=0, left=0, right=0):
    """Set cell padding in twips (1/20 of a point)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(mar)


def finalize_doc(doc, filepath):
    """Save document and print confirmation."""
    doc.save(filepath)
    print(f"  Saved: {filepath}")
    return filepath


# ============================================================
# FULL DOCUMENT BUILDER (convenience wrapper)
# ============================================================
def build_aurelian_doc(title, subtitle="", vdr_ref="", date_str="February 2026",
                       header_subtitle=""):
    """Create a complete Aurelian document skeleton ready for content.

    Returns: (doc, functions to add content)
    """
    doc = create_aurelian_doc()
    setup_styles(doc)
    add_title_page(doc, title, subtitle, vdr_ref, date_str)
    setup_header(doc, header_subtitle)
    setup_footer(doc, date_str)
    return doc

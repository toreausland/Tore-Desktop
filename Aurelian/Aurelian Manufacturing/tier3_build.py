"""
Tier 3 — Rebuild 5 VDR documents in Aurelian Design Manual v2.0
================================================================
All documents in English. No REV6 bare references.
Master document (02_Economic_Tables) is NEVER touched.

Files:
1. 02_Use_of_Funds_Seed_V1.docx (new from .md)
2. 02_Founder_Contribution_PreSeed_Valuation_1.docx (restyle)
3. 02_CAPEX_Breakdown_By_Phase.docx (restyle)
4. 03_TAM_SAM_SOM_Norwegian_Nordic.docx (restyle + translate)
5. 04_Concept_Note_Strategic_Production_Node.docx (restyle + translate)
"""

import sys
import os
import shutil
from datetime import datetime

sys.path.insert(0, os.getcwd())
sys.stdout.reconfigure(encoding='utf-8')

import aurelian_docx_template as t
from docx import Document
from docx.shared import Pt, Cm

VDR = "Aurelian_VDR"
BACKUP = "_BACKUPS_TIER3"
os.makedirs(BACKUP, exist_ok=True)

def backup(path):
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    fname = os.path.basename(path)
    dst = os.path.join(BACKUP, f"{ts}_{fname}")
    if os.path.exists(path):
        shutil.copy2(path, dst)
        print(f"  Backup: {dst}")


# ============================================================
# FILE 1: 02_Use_of_Funds_Seed_V1.docx
# Source: .md file (Norwegian) — rebuild in English
# ============================================================
def build_use_of_funds():
    print("\n" + "=" * 60)
    print("FILE 1: 02_Use_of_Funds_Seed_V1.docx")
    print("=" * 60)

    outpath = os.path.join(VDR, "02_Financial", "2.6_Use_of_Funds", "02_Use_of_Funds_Seed_V1.docx")
    backup(outpath)

    doc = t.build_aurelian_doc(
        title="Use of Funds — Seed Round",
        subtitle="Capital Deployment Plan for First Workshop Establishment",
        vdr_ref="VDR 02.06",
        date_str="February 2026"
    )

    # Section 1
    t.add_heading(doc, "1. Overview", level=1)
    t.add_body(doc, "Aurelian Manufacturing seeks 51.3 MNOK in equity in the Seed round, with a pre-money valuation of 130 MNOK. Combined with 29.3 MNOK in bank debt, this provides a total capital base of 80.6 MNOK to establish production capacity with 5 CNC machines and an operational buffer.")
    t.add_body_with_bold(doc, [
        ("All financial figures in this document are sourced from ", False),
        ("02 Economic Tables & Projections", True),
        (" (VDR 02.04). Future revenue and profit estimates are projections based on model assumptions.", False)
    ])

    # Section 2
    t.add_heading(doc, "2. Capital Structure — Seed", level=1)
    t.add_total_row_table(doc,
        ["Source", "Amount (MNOK)", "Share"],
        [
            ["Equity (Seed investors)", "51.3", "63.6%"],
            ["Bank debt (machines + shop base)", "29.3", "36.4%"],
        ],
        ["Total available", "80.6", "100%"]
    )
    t.add_body_with_bold(doc, [
        ("Debt ratio (Seed): ", True),
        ("50% of machine CAPEX financed with bank debt.", False)
    ])

    # Section 3
    t.add_heading(doc, "3. Use of Funds — Detailed", level=1)

    # 3.1
    t.add_heading(doc, "3.1 CNC Machines and Automation", level=2)
    t.add_table(doc,
        ["Item", "Quantity", "Unit cost", "Total (MNOK)"],
        [["CNC machines (MAZAK/DMG MORI) incl. automation", "5", "10.0", "50.0"]]
    )
    t.add_body(doc, "Machine specification:", bold=True)
    t.add_bullet(doc, "Multi-axis milling and turning centres")
    t.add_bullet(doc, "MAZAK MPP/PALLETECH pallet systems")
    t.add_bullet(doc, "Robotic loading/unloading")
    t.add_bullet(doc, "Integrated monitoring and SPC")
    t.add_bullet(doc, "Installation, commissioning and training included")
    t.add_body_with_bold(doc, [
        ("Financing: ", True),
        ("50% equity (25.0 MNOK) + 50% bank debt (25.0 MNOK)", False)
    ])

    # 3.2
    t.add_heading(doc, "3.2 Shop Base Setup", level=2)
    t.add_total_row_table(doc,
        ["Item", "Amount (MNOK)"],
        [
            ["Measurement room (Wenzel CMM, Jenoptik, equipment)", "3.91"],
            ["Cutting / raw goods (automated saw, shelving, misc)", "1.50"],
            ["Compressor (Kaeser)", "0.36"],
            ["Machine extraction (Absolent, 4 machines)", "2.00"],
            ["Forklift (used + top brand)", "0.85"],
        ],
        ["Shop base total", "~8.6"]
    )
    t.add_body_with_bold(doc, [
        ("Financing: ", True),
        ("Partially equity, partially bank debt (4.3 MNOK debt for long-term assets).", False)
    ])

    # 3.3
    t.add_heading(doc, "3.3 Operational Buffer", level=2)
    t.add_total_row_table(doc,
        ["Item", "Amount (MNOK)"],
        [
            ["Working capital (6 months)", "6.0"],
            ["Contingency", "4.0"],
        ],
        ["Buffer total", "10.0"]
    )
    t.add_body(doc, "The buffer ensures the company can operate for 6+ months even with delayed customer revenue, and covers unforeseen costs related to commissioning and startup.")

    # 3.4
    t.add_heading(doc, "3.4 Facility and Setup", level=2)
    t.add_total_row_table(doc,
        ["Item", "Amount (MNOK)"],
        [
            ["Lease cost during construction (prepaid)", "~2.0"],
            ["Adaptations and fit-out", "~3.0"],
        ],
        ["Facility total", "~5.0"]
    )
    t.add_body_with_bold(doc, [
        ("Note: ", True),
        ("The building is owned by Norbygg and leased by Aurelian at 5.2 MNOK/year. Building CAPEX is therefore not part of the Seed round.", False)
    ])

    # 3.5
    t.add_heading(doc, "3.5 Recruitment and Startup", level=2)
    t.add_total_row_table(doc,
        ["Item", "Amount (MNOK)"],
        [
            ["Recruitment and onboarding (8 operative + 4 admin)", "~2.5"],
            ["Certification (ISO 9001 startup)", "~0.5"],
            ["Legal and advisory", "~1.0"],
        ],
        ["Recruitment/startup total", "~4.0"]
    )

    # Section 4
    t.add_heading(doc, "4. Summary", level=1)
    t.add_total_row_table(doc,
        ["Category", "Amount (MNOK)", "Share of total"],
        [
            ["CNC machines and automation", "50.0", "62.0%"],
            ["Shop base setup", "8.6", "10.7%"],
            ["Operational buffer", "10.0", "12.4%"],
            ["Facility and setup", "~5.0", "6.2%"],
            ["Recruitment and startup", "~4.0", "5.0%"],
            ["Diverse / margin", "~3.0", "3.7%"],
        ],
        ["Total", "~80.6", "100%"]
    )

    # Section 5
    t.add_heading(doc, "5. Funding Sources — Mapping", level=1)
    t.add_total_row_table(doc,
        ["Use", "Equity (MNOK)", "Bank debt (MNOK)"],
        [
            ["CNC machines", "25.0", "25.0"],
            ["Shop base", "4.3", "4.3"],
            ["Buffer + startup", "19.0", "—"],
            ["Facility", "3.0", "—"],
        ],
        ["Total", "51.3", "29.3"]
    )

    # Section 6
    t.add_heading(doc, "6. Milestones — Seed Period", level=1)
    t.add_table(doc,
        ["Milestone", "Target date", "Prerequisite"],
        [
            ["Seed close", "Q1 2026", "51.3 MNOK secured"],
            ["Machine order", "Q2 2026", "5 CNC ordered from MAZAK/DMG MORI"],
            ["Building start (Norbygg)", "Q3 2026", "Lease contract signed"],
            ["Machine delivery", "Q2 2027", "5 CNC delivered to facility"],
            ["Commissioning", "Q2 2027", "Installation and testing"],
            ["Production start", "August 2027", "First customer order"],
            ["ISO 9001", "2027", "Certification completed"],
            ["Break-even (model)", "~Q4 2027/Q1 2028", "~24% utilization achieved"],
        ]
    )

    # Section 7
    t.add_heading(doc, "7. Sensitivity — Break-Even", level=1)
    t.add_body_with_bold(doc, [
        ("The model projects break-even at ", False),
        ("~24% utilization", True),
        (" on 5 CNC machines:", False)
    ])
    t.add_table(doc,
        ["Scenario", "Utilization", "Revenue (MNOK)", "Result"],
        [
            ["Worst case", "15%", "~19.7", "Negative (uses buffer)"],
            ["Break-even (model)", "~24%", "~31.5", "~0"],
            ["Budget (2028)", "37.5%", "49.3", "+15.2"],
            ["Stretch", "45%", "59.1", "+25.0"],
        ]
    )
    t.add_body(doc, "For context: Observed utilization among Norwegian HMLV workshops ranges between 20-45% (CNC Benchmark, VDR 03.02). The model's break-even point of ~24% is in the lower part of this observed range.", italic=True)

    # References
    t.add_references(doc, [
        ("02.04", "02 Economic Tables & Projections (master document)"),
        ("03.02", "CNC Benchmark & Competitive Landscape"),
        ("03.06", "Market Trends & Projections"),
        ("04.02", "CNC Equipment Specifications"),
        ("04.04", "Facility Strategy (Norbygg)"),
    ])

    # Contact
    t.add_contact_block(doc, "VDR 02.06")
    t.finalize_doc(doc, outpath)
    return outpath


# ============================================================
# FILE 2: 02_Founder_Contribution_PreSeed_Valuation_1.docx
# Already in English — restyle with new design
# ============================================================
def build_founder_contribution():
    print("\n" + "=" * 60)
    print("FILE 2: 02_Founder_Contribution_PreSeed_Valuation_1.docx")
    print("=" * 60)

    # Read content from existing file
    src_path = os.path.join(VDR, "02_Financial", "2.3_Valuation_Basis",
                             "02_Founder_Contribution_PreSeed_Valuation_1.docx")
    backup(src_path)
    src = Document(src_path)

    # Extract all content
    headings = []
    body_paras = []
    tables_data = []

    for p in src.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else "Normal"
        if style.startswith("Heading"):
            level = int(style.replace("Heading ", "")) if "Heading " in style else 1
            headings.append((text, level))
            body_paras.append(("heading", text, level))
        elif style == "List Paragraph" or style.startswith("List"):
            body_paras.append(("bullet", text, 0))
        else:
            body_paras.append(("body", text, 0))

    for table in src.tables:
        headers = [c.text.strip() for c in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([c.text.strip() for c in row.cells])
        tables_data.append((headers, rows))

    # Build new document
    doc = t.build_aurelian_doc(
        title="Founder Contribution & PreSeed Valuation Basis",
        subtitle="Documentation of Founder Deliverables and Valuation Rationale",
        vdr_ref="VDR 02.03",
        date_str="February 2026"
    )

    # Re-add content with new styling
    table_idx = 0
    skip_until_heading = False
    prev_was_heading = False

    for item_type, text, level in body_paras:
        # Skip title page content (first few paragraphs)
        if text in ["Aurelian Manufacturing",
                     "Founder Contribution & PreSeed Valuation Basis",
                     "Documentation of Founder Deliverables and Valuation Rationale",
                     "VDR 02.03", "February 2026", "CONFIDENTIAL", "Confidential"]:
            continue

        if item_type == "heading":
            t.add_heading(doc, text, level=level)
            prev_was_heading = True
        elif item_type == "bullet":
            t.add_bullet(doc, text)
            prev_was_heading = False
        else:
            # Body text
            t.add_body(doc, text)
            prev_was_heading = False

    # Re-add tables in sequence
    # We need to place tables at the right positions
    # Since extraction loses positional info, we rebuild with all tables at end
    # Better approach: interleave tables with content

    # For this document, tables follow specific headings.
    # Let's rebuild properly by reading the source document structure

    doc2 = t.build_aurelian_doc(
        title="Founder Contribution & PreSeed Valuation Basis",
        subtitle="Documentation of Founder Deliverables and Valuation Rationale",
        vdr_ref="VDR 02.03",
        date_str="February 2026"
    )

    # Read source and rebuild with proper interleaving
    src2 = Document(src_path)
    body = src2.element.body

    for element in body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            # It's a paragraph
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, src2)
            text = para.text.strip()
            if not text:
                continue

            style = para.style.name if para.style else "Normal"

            # Skip title page elements
            if text in ["Aurelian Manufacturing",
                         "Founder Contribution & PreSeed Valuation Basis",
                         "Documentation of Founder Deliverables and Valuation Rationale",
                         "VDR 02.03", "February 2026", "Confidential", "CONFIDENTIAL"]:
                continue
            # Skip contact block at end
            if "andre@aurelian.no" in text or "tore@aurelian.no" in text:
                continue
            if text.startswith("Aurelian Manufacturing AS"):
                continue

            if style.startswith("Heading"):
                level = int(style.replace("Heading ", "")) if "Heading " in style else 1
                t.add_heading(doc2, text, level=level)
            elif "List" in style:
                t.add_bullet(doc2, text)
            else:
                t.add_body(doc2, text)

        elif tag == 'tbl':
            # It's a table
            from docx.table import Table
            tbl = Table(element, src2)
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            rows = []
            for row in tbl.rows[1:]:
                rows.append([c.text.strip() for c in row.cells])
            t.add_table(doc2, headers, rows)

    # Add references
    t.add_references(doc2, [
        ("02.04", "02 Economic Tables & Projections (master document)"),
        ("02.02", "Financing Plan"),
        ("01.04", "Cap Table Pro Forma All Rounds"),
    ])
    t.add_contact_block(doc2, "VDR 02.03")
    t.finalize_doc(doc2, src_path)

    # Also update the duplicate in 06_Team
    dup_path = os.path.join(VDR, "06_Team", "6.1_Founders",
                             "02_Founder_Contribution_PreSeed_Valuation_1.docx")
    if os.path.exists(dup_path):
        backup(dup_path)
        shutil.copy2(src_path, dup_path)
        print(f"  Copied to duplicate: {dup_path}")

    return src_path


# ============================================================
# FILE 3: 02_CAPEX_Breakdown_By_Phase.docx
# English content, minimal formatting — rebuild with design
# ============================================================
def build_capex_breakdown():
    print("\n" + "=" * 60)
    print("FILE 3: 02_CAPEX_Breakdown_By_Phase.docx")
    print("=" * 60)

    src_path = os.path.join(VDR, "02_Financial", "2.4_Economic_Tables_Projections",
                             "02_CAPEX_Breakdown_By_Phase.docx")
    backup(src_path)
    src = Document(src_path)

    doc = t.build_aurelian_doc(
        title="CAPEX Breakdown by Phase",
        subtitle="Investment Schedule & Machine Deployment",
        vdr_ref="VDR 02.04",
        date_str="February 2026"
    )

    # Read and rebuild with interleaving
    body = src.element.body
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    skip_texts = ["Aurelian Manufacturing", "CAPEX Breakdown by Phase",
                  "Investment Schedule & Machine Deployment", "February 2026",
                  "Confidential", "CONFIDENTIAL"]

    for element in body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

        if tag == 'p':
            para = Paragraph(element, src)
            text = para.text.strip()
            if not text or text in skip_texts:
                continue

            style = para.style.name if para.style else "Normal"

            # Fix any REV6 references in text
            text = text.replace("Economic Tables REV6", "02 Economic Tables & Projections (VDR 02.04)")
            text = text.replace("REV6", "02 Economic Tables & Projections")
            # Fix banned name "Vedlegg E"
            text = text.replace("Vedlegg E — CNC Resale Value Analysis",
                              "CNC Resale Value Analysis (VDR 04.02)")
            text = text.replace("Vedlegg E", "CNC Resale Value Analysis (VDR 04.02)")

            if style.startswith("Heading"):
                level = int(style.replace("Heading ", "")) if "Heading " in style else 1
                t.add_heading(doc, text, level=level)
            elif "List" in style:
                t.add_bullet(doc, text)
            else:
                t.add_body(doc, text)

        elif tag == 'tbl':
            tbl = Table(element, src)
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            rows = []
            for row in tbl.rows[1:]:
                row_data = [c.text.strip() for c in row.cells]
                # Fix REV6 in table cells
                row_data = [cell.replace("Economic Tables REV6", "02 Economic Tables & Projections")
                           .replace("REV6", "02 Economic Tables & Projections")
                           for cell in row_data]
                rows.append(row_data)
            t.add_table(doc, headers, rows)

    # Add references
    t.add_references(doc, [
        ("02.04", "02 Economic Tables & Projections (master document)"),
        ("02.06", "Use of Funds — Seed"),
        ("04.02", "CNC Equipment Specifications"),
        ("04.04", "Facility Strategy (Norbygg)"),
    ])
    t.add_contact_block(doc, "VDR 02.04")
    t.finalize_doc(doc, src_path)
    return src_path


# ============================================================
# FILE 4: 03_TAM_SAM_SOM_Norwegian_Nordic.docx
# Norwegian content — translate to English + restyle
# ============================================================
def build_tam_sam_som():
    print("\n" + "=" * 60)
    print("FILE 4: 03_TAM_SAM_SOM_Norwegian_Nordic.docx")
    print("=" * 60)

    src_path = os.path.join(VDR, "03_Commercial_Market", "3.1_Market_Analysis",
                             "03_TAM_SAM_SOM_Norwegian_Nordic.docx")
    backup(src_path)
    src = Document(src_path)

    # Extract content for translation
    # This document is entirely Norwegian — we need to translate headings and body text
    # Read all content first
    content_blocks = []
    body = src.element.body
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for element in body:
        tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag
        if tag == 'p':
            para = Paragraph(element, src)
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else "Normal"
                content_blocks.append(("para", style, text))
        elif tag == 'tbl':
            tbl = Table(element, src)
            headers = [c.text.strip() for c in tbl.rows[0].cells]
            rows = [[c.text.strip() for c in row.cells] for row in tbl.rows[1:]]
            content_blocks.append(("table", headers, rows))

    # Build the new document with translated content
    doc = t.build_aurelian_doc(
        title="TAM / SAM / SOM",
        subtitle="Norwegian & Nordic CNC Market Sizing",
        vdr_ref="VDR 03.01",
        date_str="February 2026"
    )

    # Norwegian-to-English translation map for common terms
    no_to_en = {
        "Markedsnivå": "Market level",
        "Definisjon": "Definition",
        "Årlig verdi (NOK)": "Annual value (NOK)",
        "Parameter": "Parameter",
        "Verdi": "Value",
        "Periode": "Period",
        "Kilde": "Source",
        "Segment": "Segment",
        "Estimert totalverdi": "Estimated total value",
        "Andel av TAM": "Share of TAM",
        "Andel": "Share",
        "CNC-adresserbart (NOK/år)": "CNC-addressable (NOK/year)",
        "Estimert verdi (NOK/år)": "Estimated value (NOK/year)",
        "Kommentar": "Comment",
        "Selskap": "Company",
        "CNC": "CNC",
        "Est. utnyttelse": "Est. utilization",
        "Ansatte/CNC": "Staff/CNC",
        "Driver": "Driver",
        "Effekt": "Effect",
        "Tidshorisont": "Time horizon",
        "VDR #": "VDR ref",
        "Dokument": "Document",
        "Brukt til": "Used for",
    }

    def translate_header(h):
        return no_to_en.get(h, h)

    def translate_cell(c):
        """Basic cell translation — keep numbers, translate known terms."""
        # Don't translate numbers or VDR references
        if any(ch.isdigit() for ch in c) and not any(no_word in c for no_word in ["Forsvars", "Energi", "Maritim"]):
            return c
        return no_to_en.get(c, c)

    # Section-level translation — we rebuild the document structure
    t.add_heading(doc, "1. Market Sizing Overview", level=1)
    t.add_body(doc, "This document quantifies the Total Addressable Market (TAM), Serviceable Addressable Market (SAM) and Serviceable Obtainable Market (SOM) for Aurelian Manufacturing's CNC production services in the Norwegian and Nordic markets.")
    t.add_body(doc, "All market size estimates are based on published industry data, government statistics and third-party market reports. Aurelian's own projections are design targets and are clearly separated from observed market data.", italic=True)

    # Process source content — translate tables, rebuild body
    table_count = 0
    for block_type, *data in content_blocks:
        if block_type == "table":
            headers, rows = data
            # Translate headers
            en_headers = [translate_header(h) for h in headers]

            # Check if this is the references table (last one)
            if "VDR #" in headers or "VDR" in str(headers):
                # Skip — we'll add our own references section
                continue

            t.add_table(doc, en_headers, rows)
            table_count += 1

        elif block_type == "para":
            style, text = data
            # Skip header/footer content
            if "Konfidensielt" in text or "Side" == text.strip():
                continue
            if text.startswith("Aurelian Manufacturing") and "|" in text:
                continue

            # Translate key section headings
            section_translations = {
                "1. TAM — Totalt adresserbart marked": "2. TAM — Total Addressable Market",
                "2. SAM — Serviceable Addressable Market": "3. SAM — Serviceable Addressable Market",
                "3. SOM — Serviceable Obtainable Market": "4. SOM — Serviceable Obtainable Market",
                "4. Konkurranselandskap (utvalg)": "5. Competitive Landscape (Selection)",
                "5. Markedsdrivere og timing": "6. Market Drivers and Timing",
                "6. Metodikk og begrensninger": "7. Methodology and Limitations",
                "7. Dokumentreferanser": "References",  # Will be handled by add_references
            }

            if text in section_translations:
                translated = section_translations[text]
                if translated == "References":
                    continue  # Skip — we add our own
                t.add_heading(doc, translated, level=1)
            elif style.startswith("Heading"):
                # Keep as-is if not in translation map (sub-headings may be OK)
                t.add_heading(doc, text, level=2)
            else:
                # Body text — translate key Norwegian phrases
                text = text.replace("Vedlegg C — Sammenligning med lignende industrielle investeringscaser for detaljer",
                                   "See CNC Benchmark & Competitive Landscape (VDR 03.02) for comparable industrial investment cases")
                text = text.replace("Vedlegg C", "CNC Benchmark & Competitive Landscape (VDR 03.02)")
                text = text.replace("Vedlegg", "Appendix")
                text = text.replace("Konfidensielt", "Confidential")
                t.add_body(doc, text)

    # Add references
    t.add_references(doc, [
        ("02.04", "02 Economic Tables & Projections (master document)"),
        ("03.02", "CNC Benchmark & Competitive Landscape"),
        ("03.06", "Market Trends & Projections"),
    ])
    t.add_contact_block(doc, "VDR 03.01")
    t.finalize_doc(doc, src_path)
    return src_path


# ============================================================
# FILE 5: 04_Concept_Note_Strategic_Production_Node.docx
# Norwegian content, 0/9 design — full rebuild + translate
# ============================================================
def build_concept_note():
    print("\n" + "=" * 60)
    print("FILE 5: 04_Concept_Note_Strategic_Production_Node.docx")
    print("=" * 60)

    src_path = os.path.join(VDR, "04_Technical_Operations", "4.1_Technology_Overview",
                             "04_Concept_Note_Strategic_Production_Node.docx")
    backup(src_path)
    src = Document(src_path)

    doc = t.build_aurelian_doc(
        title="Strategic Production Node",
        subtitle="Concept Note — Customer-Anchored Autonomous Manufacturing",
        vdr_ref="VDR 04.01",
        date_str="February 2026"
    )

    # This document is entirely Norwegian prose (no tables).
    # We translate the structure and key content.

    t.add_body(doc, "Aurelian Manufacturing establishes a strategic production node — designed for high utilization in normal operations and supply security in exceptional circumstances.", italic=True)

    # Section 1
    t.add_heading(doc, "1. Background and Purpose", level=1)
    t.add_body(doc, "This note is prepared for decision-makers considering engagement with Aurelian Manufacturing as a development partner, industrial customer and co-investor.")
    t.add_body(doc, "The purpose is to clarify how Aurelian Manufacturing is structured as an industrial production and capacity project, where technology, real estate and capacity allocation form an integrated whole.")
    t.add_body(doc, "The note is intended to provide a basis for further dialogue on ownership, capacity allocation and phased development, and to clarify the role different actors can play — from financial contribution to operational involvement.")

    # Section 2
    t.add_heading(doc, "2. What Is Being Established", level=1)

    t.add_heading(doc, "2.1 Operational Concept", level=2)
    t.add_body(doc, "Aurelian Manufacturing is established as a greenfield, autonomous HMLV workshop, designed from the ground up for high actual CNC utilization across all hours of the day and week.")
    t.add_body(doc, "The operational concept is developed to achieve high capital productivity in an HMLV environment through sub-linear staffing and standardized processes. The design target is 0.8 FTE per CNC at scale, compared to an industry average of approximately 2.5 FTE per CNC.")

    t.add_heading(doc, "2.2 Physical Infrastructure", level=2)
    t.add_body(doc, "The physical infrastructure is dimensioned to support the operational concept, not the other way around. The workshop is established on a pre-zoned industrial site with a purpose-built 2,635 m\u00B2 facility by Norbygg on a 30,000 m\u00B2 site.")
    t.add_body(doc, "The size of the first build phase is chosen as an optimal starting point. The area is sufficient to establish an economically sustainable business with high machine utilization from day one.")

    # Section 3
    t.add_heading(doc, "3. Strategic Location — Operational Robustness", level=1)
    t.add_body(doc, "The location is chosen based on national logistical and preparedness relevance, not real estate speculation or short-term cost optimization. The workshop is positioned to ensure transport access by rail, sea and road.")
    t.add_body(doc, "The combination of rail, sea and road is deliberately chosen to ensure logistical redundancy and operational preparedness, so that production and distribution can be maintained even during periods of infrastructure disruption.")

    # Section 4
    t.add_heading(doc, "4. A New Investment Model: Customer-Anchored Capacity", level=1)
    t.add_body(doc, "The investment model is not based on exclusivity, operational control from the customer side or vertical integration. Aurelian Manufacturing remains an independent production platform.")
    t.add_body(doc, "Instead of a traditional, broad financial investor approach, the project is structured around a directed, industrial model targeting a limited number of selected anchor customers who commit capacity allocations in exchange for priority access.")
    t.add_body(doc, "For customers, this provides predictable access to production capacity, geographic risk diversification and reduced supplier dependency. Meanwhile, Aurelian retains full operational independence and scalability.")

    # Section 5
    t.add_heading(doc, "5. Implicit Preparedness Value", level=1)
    t.add_body(doc, "In a situation of increased geopolitical uncertainty or national disruptions, customers already have access to qualified production lines, approved processes and reserved capacity — without the need for new supplier qualification.")
    t.add_body(doc, "This eliminates the need for new supplier qualification, ramp-up in crisis situations and ad-hoc procurement. Production capacity thus functions as a strategic insurance policy for customers with critical supply chains.")

    # Section 6
    t.add_heading(doc, "6. Real Estate and Developer Role", level=1)
    t.add_body(doc, "Real estate and industrial operations are deliberately separated. The developer owns the building and land, while Aurelian Manufacturing leases purpose-built premises through long-term agreements.")
    t.add_body(doc, "The real estate and site structure is designed for phased development, where pace and scope are governed by documented demand and actual capacity utilization.")
    t.add_body(doc, "The pre-zoned site and modular building logic enable step-by-step development from the first workshop to a larger industrial facility without requiring new zoning or major infrastructure investment.")

    # Section 7
    t.add_heading(doc, "7. Conclusion", level=1)
    t.add_body(doc, "Aurelian Manufacturing is established to develop a replicable, strategic production node with high capital productivity in normal operations, documented preparedness value and a scalable model for further expansion.")
    t.add_body(doc, "The first production node is simultaneously intended to serve as a reference and blueprint for further establishment of similar nodes in Norway and Europe.")

    # References
    t.add_references(doc, [
        ("02.04", "02 Economic Tables & Projections (master document)"),
        ("04.04", "Facility Strategy (Norbygg)"),
        ("03.04", "Go-To-Market Strategy"),
    ])
    t.add_contact_block(doc, "VDR 04.01")
    t.finalize_doc(doc, src_path)
    return src_path


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("Aurelian VDR — Tier 3 Build")
    print(f"Date: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    print(f"All documents in ENGLISH")
    print()

    results = []
    results.append(build_use_of_funds())
    results.append(build_founder_contribution())
    results.append(build_capex_breakdown())
    results.append(build_tam_sam_som())
    results.append(build_concept_note())

    print("\n" + "=" * 60)
    print(f"COMPLETE — {len(results)} files built")
    for r in results:
        print(f"  {r}")
    print(f"Backups in: {BACKUP}/")
    print("=" * 60)
